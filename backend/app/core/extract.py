from pypdf import PdfReader
import pypdfium2
import pytesseract
import docx

MIN_CHARS = 50

def render_page_to_image(filepath: str, page_number: int):
    pdf = pypdfium2.PdfDocument(filepath)
    page = pdf[page_number]
    bitmap = page.render(scale=2.0)
    return bitmap.to_pil()

def ocr_page(filepath: str, page_number: int) -> str:
    image = render_page_to_image(filepath, page_number)
    return pytesseract.image_to_string(image)

def ocr_embedded_images(page) -> list[str]:
    results = []
    for img in page.images:
        pil_image = img.image
        text = pytesseract.image_to_string(pil_image)
        if text.strip():
            results.append(text)
    return results


def extract_document(filepath: str) -> list[dict]:
    if filepath.endswith(".pdf"):
        reader = PdfReader(filepath)
        pages = []
        for i, page in enumerate (reader.pages):
            text = page.extract_text().strip()
            images = page.images

            if not images:
                final_text = text if len(text) >= MIN_CHARS else ocr_page(filepath, i)
            
            elif len(text) < MIN_CHARS:
                final_text = ocr_page(filepath, i)

            else:
                image_text = ocr_embedded_images(page)
                final_text = text
                for img_text in image_text:
                    pages.append({
                        "page_number": i + 1,
                        "content": img_text
                    })

            pages.append({
                "page_number": i + 1,
                "content": final_text
            })

        return pages
    
    elif filepath.endswith(".docx"):
        doc = docx.Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        return [{"page_number": 1, "content": text}]

    elif filepath.endswith((".txt", ".md")):
        with open(filepath, "r", encoding="UTF-8") as f:
            return [{"page_number": 1, "content": f.read()}]
        
    else:
        raise ValueError(f"Unsupported file type: {filepath}")